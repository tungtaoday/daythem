"""Gemini (Google AI Studio) adapter — trích danh sách tên học sinh từ FILE bất kỳ.

Hỗ trợ: ảnh (jpg/png/webp/heic), PDF, DOCX, XLSX, CSV/TXT.
- Ảnh & PDF: đưa thẳng Gemini (vision/document).
- DOCX/XLSX: bóc text ở server (python-docx/openpyxl) rồi mới đưa Gemini
  (Gemini KHÔNG đọc trực tiếp định dạng Office nhị phân).

Khoá API ở server; app không thấy khoá. Không lưu file — chỉ trích tên rồi bỏ.
"""
import io
import json
import base64
import urllib.request
import urllib.error
from daythem.config import settings

_ENDPOINT = "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"

_INSTRUCTION = (
    "Trích xuất HỌ TÊN ĐẦY ĐỦ của từng học sinh trong danh sách, giữ nguyên dấu tiếng Việt. "
    "Bỏ số thứ tự, tiêu đề, tên lớp, tên môn, ngày tháng, số điện thoại, và mọi thông tin không phải tên học sinh. "
    "Nếu không phải danh sách tên thì trả về mảng rỗng."
)

# Chặn file quá lớn (base64). ~10MB base64 ≈ ~7.5MB file.
_MAX_B64_LEN = 10_000_000
# Gemini nhận trực tiếp media này (không cần bóc text):
_MEDIA_MIMES = ("image/", "application/pdf")


def _call_gemini(parts: list[dict]) -> list[str]:
    body = {
        "contents": [{"parts": parts}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": {"type": "ARRAY", "items": {"type": "STRING"}},
            "temperature": 0,
        },
    }
    url = _ENDPOINT.format(model=settings.GEMINI_MODEL, key=settings.GEMINI_API_KEY)
    req = urllib.request.Request(
        url, data=json.dumps(body).encode(), headers={"Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = json.load(resp)
    except (urllib.error.URLError, TimeoutError, OSError) as e:
        raise ValueError("Không đọc được file lúc này, thử lại sau.") from e
    try:
        text = data["candidates"][0]["content"]["parts"][0]["text"]
        names = json.loads(text)
    except (KeyError, IndexError, TypeError, json.JSONDecodeError):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for n in names if isinstance(names, list) else []:
        if not isinstance(n, str):
            continue
        n = " ".join(n.split())
        if 1 < len(n) <= 60 and n.lower() not in seen:
            seen.add(n.lower())
            out.append(n)
    return out[:200]


def _names_from_media(b64: str, mime: str) -> list[str]:
    return _call_gemini([
        {"text": "Đây là ảnh/tài liệu danh sách học sinh. " + _INSTRUCTION},
        {"inline_data": {"mime_type": mime, "data": b64}},
    ])


def _names_from_text(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    return _call_gemini([
        {"text": "Dưới đây là nội dung trích từ một file danh sách học sinh. " + _INSTRUCTION
                 + "\n\n--- NỘI DUNG ---\n" + text[:20000]},
    ])


def _docx_to_text(raw: bytes) -> str:
    import docx  # python-docx
    doc = docx.Document(io.BytesIO(raw))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text for c in row.cells))
    return "\n".join(l for l in lines if l.strip())


def _xlsx_to_text(raw: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _kind(mime: str, filename: str) -> str:
    m = (mime or "").lower()
    fn = (filename or "").lower()
    if m.startswith("image/") or fn.endswith((".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif")):
        return "media"
    if m == "application/pdf" or fn.endswith(".pdf"):
        return "pdf"
    if "wordprocessingml" in m or fn.endswith(".docx"):
        return "docx"
    if "spreadsheetml" in m or fn.endswith((".xlsx", ".xlsm")):
        return "xlsx"
    if m in ("text/csv", "text/plain") or fn.endswith((".csv", ".txt")):
        return "text"
    return "unknown"


def extract_student_names(file_b64: str, mime_type: str = "", filename: str = "") -> list[str]:
    """Trích tên học sinh từ file (ảnh/pdf/docx/xlsx/csv/txt). Raise ValueError khi lỗi/không hỗ trợ."""
    if not settings.GEMINI_API_KEY:
        raise ValueError("Tính năng nhập từ file chưa được bật.")
    if not file_b64 or len(file_b64) > _MAX_B64_LEN:
        raise ValueError("File không hợp lệ hoặc quá lớn.")

    kind = _kind(mime_type, filename)
    if kind == "media":
        return _names_from_media(file_b64, mime_type or "image/jpeg")
    if kind == "pdf":
        return _names_from_media(file_b64, "application/pdf")

    try:
        raw = base64.b64decode(file_b64)
    except Exception as e:
        raise ValueError("Không giải mã được file.") from e

    if kind == "docx":
        try:
            return _names_from_text(_docx_to_text(raw))
        except ValueError:
            raise
        except Exception as e:
            raise ValueError("Không đọc được file Word.") from e
    if kind == "xlsx":
        try:
            return _names_from_text(_xlsx_to_text(raw))
        except ValueError:
            raise
        except Exception as e:
            raise ValueError("Không đọc được file Excel.") from e
    if kind == "text":
        return _names_from_text(raw.decode("utf-8", errors="ignore"))

    raise ValueError("Định dạng file chưa hỗ trợ. Hãy dùng ảnh, PDF, Word (.docx), Excel (.xlsx) hoặc CSV.")
