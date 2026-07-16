"""Sinh file DOCX tờ khai thuế (bản nháp tham khảo) cho hộ/cá nhân kinh doanh dạy thêm.

Mẫu: 01/TKN-CNKD (doanh thu ≤ ngưỡng) hoặc 01/CNKD (> ngưỡng). Đây là BẢN NHÁP hỗ trợ,
KHÔNG phải mẫu điện tử chính thức — người nộp phải đối chiếu mẫu gốc mới nhất tại cơ quan thuế.
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _vnd(v: float) -> str:
    return f"{int(round(v)):,}".replace(",", ".") + " đ"


def _center_bold(doc: Document, text: str, size: int = 11, italic: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = not italic
    r.italic = italic
    r.font.size = Pt(size)
    return p


def build_declaration_docx(decl: dict) -> bytes:
    """Nhận dict trả về từ handle_get_tax_declaration, dựng file .docx trả về bytes."""
    f = decl["fields"]
    doc = Document()

    _center_bold(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 11)
    _center_bold(doc, "Độc lập - Tự do - Hạnh phúc", 11)
    doc.add_paragraph()

    _center_bold(doc, "TỜ KHAI THUẾ", 15)
    _center_bold(doc, f"(Hộ kinh doanh, cá nhân kinh doanh — Mẫu số {f['form_code']})", 10, italic=True)
    _center_bold(
        doc,
        "Giáo viên dạy thêm đã đăng ký hộ kinh doanh theo Thông tư 29/2024",
        9, italic=True,
    )
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.add_run(f"Kỳ tính thuế: ").bold = True
    info.add_run(f"Năm {f['year']}     Kỳ khai: {f['ky_khai']}")

    for label, value in [
        ("Người nộp thuế", f["full_name"] or "..."),
        ("Mã số thuế", f["mst"] or "..."),
        ("CCCD/CMND", f["id_number"] or "..."),
        ("Ngày sinh", f["date_of_birth"] or "..."),
        ("Địa chỉ", f["address"] or "..."),
        ("Ngành nghề", "Dịch vụ dạy học (giáo dục)"),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(str(value))

    doc.add_paragraph()

    # Bảng tính thuế
    rows = [
        ("Chỉ tiêu", "Số tiền / Tỷ lệ"),
        ("Tổng doanh thu năm", _vnd(f["tong_thu_nhap"])),
        ("Ngưỡng doanh thu miễn thuế (NĐ 141/2026)", _vnd(f["nguong_mien_thue"])),
        ("Thuế GTGT (dạy học KHÔNG chịu thuế GTGT)", "0 đ"),
        ("Doanh thu tính thuế TNCN", _vnd(f["thu_nhap_chiu_thue"])),
        ("Tỷ lệ thuế TNCN", f"{f['thue_suat'] * 100:.0f}%"),
        ("Số thuế TNCN phải nộp", _vnd(f["so_thue_phai_nop"])),
        ("TỔNG SỐ THUẾ PHẢI NỘP", _vnd(f["so_thue_phai_nop"])),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = a
        cells[1].text = b
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i == 0 or i == len(rows) - 1:
            for c in cells:
                for para in c.paragraphs:
                    for r in para.runs:
                        r.bold = True

    doc.add_paragraph()
    st = doc.add_paragraph()
    st.add_run(f["status_text"]).bold = True
    dl = doc.add_paragraph()
    dl.add_run("Hạn nộp tờ khai: ").bold = True
    dl.add_run(f["deadline"])

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("......, ngày ...... tháng ...... năm ......").italic = True
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    s2.add_run("NGƯỜI NỘP THUẾ").bold = True
    s3 = doc.add_paragraph()
    s3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    s3.add_run("(Ký, ghi rõ họ tên)").italic = True

    doc.add_paragraph()
    warn_title = doc.add_paragraph()
    warn_title.add_run("⚠ LƯU Ý QUAN TRỌNG").bold = True
    for line in [
        "Đây là BẢN NHÁP tham khảo, không thay thế tư vấn của cơ quan thuế/kế toán và không phải mẫu điện tử chính thức.",
        "Chính sách thuế hộ kinh doanh thay đổi liên tục năm 2026 (NĐ 68/2026, 141/2026; TT 18/2026, 50/2026). "
        "Hãy tải mẫu gốc và đối chiếu ngưỡng mới nhất tại gdt.gov.vn hoặc chi cục thuế trước khi nộp.",
        "Thuế khoán đã bãi bỏ từ 01/01/2026 — hộ kinh doanh khai theo phương pháp KÊ KHAI.",
        "Mẫu này áp dụng khi bạn ĐÃ đăng ký hộ kinh doanh. Nếu chỉ nhận tiền công từ trung tâm (làm thuê) "
        "thì đó là thu nhập tiền lương/tiền công — cơ chế khác, không dùng mẫu này.",
        "GieoChữ chỉ hỗ trợ tính toán và điền nháp; việc ký và nộp do người nộp thuế tự thực hiện.",
    ]:
        wp = doc.add_paragraph(line, style=None)
        wp.paragraph_format.left_indent = Pt(10)
        for r in wp.runs:
            r.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
